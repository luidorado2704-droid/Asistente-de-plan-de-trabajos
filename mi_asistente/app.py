# app.py
import streamlit as st
import requests  # Para llamar a la API de Semantic Scholar y a Groq
from docx import Document  # Para generar el Word
from docx.shared import Pt
import io  # Para crear el archivo en memoria

# --- Configuración de APIs ---
# Configura la API Key de Groq (¡Agrégalo a tus secrets.toml!)
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"

# --- Función para llamar a Groq ---
def llamar_groq(mensajes, model="llama-3.3-70b-versatile"):
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": model,
        "messages": mensajes,
        "temperature": 0.1  # Baja temperatura para respuestas más precisas
    }
    try:
        response = requests.post(GROQ_API_URL, json=data, headers=headers)
        response.raise_for_status()  # Lanza error si la solicitud falla
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"Error al llamar a Groq: {e}")
        return None

# --- Configuración de la Página ---
st.set_page_config(page_title="Asistente de Diseño Experimental", page_icon="🧪", layout="wide")
st.title("🧪 Asistente para Planes de Investigación en Biotecnología")
st.write("""
Esta herramienta ayuda a transformar una idea de investigación en un plan de trabajo detallado y basado en la literatura científica.
Ingrese su idea en el cuadro de texto y haga clic en 'Generar Plan'.
""")

# --- Widget de Entrada ---
idea_del_usuario = st.text_area(
    label="**Describa su idea de investigación en biotecnología:**",
    height=100,
    placeholder="Ej: Evaluar capacidad antioxidante y antiinflamatoria de hidrolizados proteicos de mistol y chañar..."
)

boton_generar = st.button("🚀 Generar Plan de Trabajo Basado en Literatura", type="primary")

# --- Función para Buscar en Semantic Scholar ---
def buscar_papers(terminos_busqueda):
    """Busca papers en Semantic Scholar y devuelve una lista de abstracts."""
    abstracts = []
    titulos = []
    url = "https://api.semanticscholar.org/graph/v1/paper/search"
    params = {
        'query': terminos_busqueda,
        'limit': 10,  # Número de resultados
        'fields': 'title,abstract'  # Campos a recuperar
    }
    
    try:
        response = requests.get(url, params=params)
        response.raise_for_status()  # Lanza error si la solicitud falla
        datos = response.json()
        
        for paper in datos.get('data', []):
            if paper.get('abstract'):
                abstracts.append(paper['abstract'])
                titulos.append(paper.get('title', 'Sin título'))
        return abstracts, titulos
    except requests.exceptions.RequestException as e:
        st.error(f"Error al buscar en Semantic Scholar: {e}")
        return [], []

# --- Función para crear el documento Word ---
def crear_documento_word(contenido, titulo="Plan_de_Investigacion"):
    """Crea un documento Word (.docx) a partir del texto generado."""
    doc = Document()
    # Añade un título al documento
    titulo_doc = doc.add_heading(titulo, level=0)
    titulo_doc.alignment = 1  # Centrado

    # Divide el contenido en secciones (suponiendo que está en Markdown)
    secciones = contenido.split('**')
    for i, seccion in enumerate(secciones):
        if seccion.strip() and i % 2 == 1:  # Es un título de sección (negrita en Markdown)
            heading = doc.add_heading(seccion.strip(), level=1)
        else:
            # Es texto normal
            p = doc.add_paragraph(seccion.strip())
    # Guarda el documento en un buffer de memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Lógica Principal al Presionar el Botón ---
if boton_generar:
    if idea_del_usuario:
        with st.spinner('Fase 1/3: Generando términos de búsqueda optimizados...'):
            # --- PASO 1: Prompt para generar términos de búsqueda ---
            prompt_terminos = f"""
            Eres un experto en diseño experimental en biotecnología y búsqueda de literatura científica.
            TU TAREA: Dada la siguiente idea de investigación, genera una lista de TÉRMINOS DE BÚSQUEDA OPTIMIZADOS para usar en una API de artículos científicos (como Semantic Scholar).
            Los términos deben ser precisos, relevantes y capturar los conceptos clave de la idea. Devuélvelos como una sola cadena de texto separada por comas.

            IDEA: {idea_del_usuario}

            TÉRMINOS DE BÚSQUEDA:
            """
            # LLAMADA A GROQ (reemplazo de Gemini)
            mensajes_terminos = [{"role": "user", "content": prompt_terminos}]
            response_terminos = llamar_groq(mensajes_terminos)
            
            if response_terminos is None:
                st.error("Error al generar términos de búsqueda. Intenta nuevamente.")
                st.stop()
                
            terminos_busqueda = response_terminos.strip()
            st.write("**Términos de búsqueda generados:**", terminos_busqueda)

        with st.spinner('Fase 2/3: Buscando los artículos científicos más relevantes...'):
            # --- PASO 2: Búsqueda en Semantic Scholar ---
            abstracts, titulos = buscar_papers(terminos_busqueda)
            if not abstracts:
                st.warning("No se encontraron artículos con abstracts. Generando plan basado solo en conocimiento general.")
                contexto_literatura = "No se encontraron artículos recientes específicos."
            else:
                # Une los primeros 5-7 abstracts para no exceder el contexto del modelo
                contexto_literatura = "\n\n".join([f"Título: {titulo}\nAbstract: {abstract}" for titulo, abstract in zip(titulos[:7], abstracts[:7])])

        with st.spinner('Fase 3/3: Generando el plan de trabajo con base en la literatura...'):
            # --- PASO 3: Prompt FINAL con contexto de la literatura ---
            prompt_final = f"""
            Eres un experto en diseño experimental en biotecnología. Tu tarea es generar un plan de trabajo detallado y factible para la idea de investigación proporcionada.

            **IDEA DEL INVESTIGADOR:**
            {idea_del_usuario}

            **CONTEXTO DE LITERATURA CIENTÍFICA RECIENTE (Abstracts de artículos relevantes):**
            {contexto_literatura}

            **INSTRUCCIONES PARA EL PLAN:**
            Basándote en la idea y en la literatura científica proporcionada, genera un plan de trabajo que incluya las siguientes secciones en español:

            1.  **Título Tentativo del Proyecto:** [Sugiere un título basado en la idea y la literatura]
            2.  **Introducción y Estado del Arte:** [Síntesis breve del contexto, citando hallazgos clave de los abstracts si son relevantes]
            3.  **Objetivos:**
                - **Objetivo General:** [Uno]
                - **Objetivos Específicos:** [3-5]
            4.  **Metodología Propuesta (DETALLADA):**
                - **Diseño Experimental:** [Tipo de estudio, grupos, réplicas]
                - **Extracción de Proteínas:** [Protocolo sugerido, equipos]
                - **Hidrólisis Enzimática:** [Enzimas a probar, condiciones]
                - **Ensayo de Capacidad Antioxidante:** [Métodos específicos, e.g., ORAC, DPPH, ABTS]
                - **Ensayo de Actividad Antiinflamatoria:** [Métodos específicos, e.g., inhibición de COX-2, ensayos con células]
                - **Análisis Estadístico:** [Software, tests a usar]
            5.  **Cronograma Tentativo:** [Dividido por trimestres o meses, con las actividades principales]
            6.  **Recursos Necesarios:** [Reactivos, equipos, software, personal]

            **Formato:** Usa **negritas** para los títulos de sección (ej: **4. Metodología Propuesta**) y listas con viñetas para los items.
            """
            
            # LLAMADA A GROQ (reemplazo de Gemini)
            mensajes_final = [{"role": "user", "content": prompt_final}]
            response_final = llamar_groq(mensajes_final)
            
            if response_final is None:
                st.error("Error al generar el plan final. Intenta nuevamente.")
                st.stop()
                
            plan_final = response_final

        # --- MOSTRAR RESULTADO ---
        st.success("¡Plan generado con éxito! Basado en el análisis de la literatura científica.")
        st.write(plan_final)

        # --- BOTÓN DE DESCARGA PARA WORD ---
        doc_buffer = crear_documento_word(plan_final, titulo=f"Plan_{idea_del_usuario[:20]}")
        st.download_button(
            label="📥 Descargar Plan en Word (.docx)",
            data=doc_buffer,
            file_name="plan_investigacion.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        st.warning("Por favor, ingrese una idea de investigación primero.")